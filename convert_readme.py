import markdown
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def markdown_to_docx(md_file, docx_file):
    # Read Markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    doc = Document()
    
    # Title
    doc.add_heading('Pronto - Sustainable BI Platform', 0)

    # Simple parsing logic (robust enough for README structure)
    lines = md_content.split('\n')
    
    code_mode = False
    
    for line in lines:
        line = line.strip()
        
        # Code blocks
        if line.startswith('```'):
            code_mode = not code_mode
            continue
            
        if code_mode:
            p = doc.add_paragraph()
            p.style = 'Quote' 
            run = p.add_run(line)
            run.font.name = 'Courier New'
            continue
            
        if not line:
            continue
            
        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
            
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\.', line):
            # Ordered list
            text = re.sub(r'^\d+\.\s*', '', line)
            p = doc.add_paragraph(text, style='List Number')
            
        # Normal text
        else:
            doc.add_paragraph(line)

    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == "__main__":
    markdown_to_docx("README.md", "Pronto_README.docx")
